import pandas as pd
import json
import re
from docx import Document
from pathlib import Path
import os
import sys

# Fix Unicode encoding for Windows console
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# Define paths
DATA_DIR = Path('data')
OUTPUT_DIR = Path('output')

# Ensure directories exist
DATA_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

def extract_faculty_from_excel(filepath):
    """Extract faculty list and course assignments from Excel file"""
    try:
        df = pd.read_excel(filepath, sheet_name='Subject allotment', header=None)
        
        faculty_data = []
        course_assignments = []
        
        in_table = False
        for idx, row in df.iterrows():
            if idx > 200:
                break
            
            if idx == 9:
                in_table = True
                continue
            
            if in_table and idx > 9 and idx < 50:
                if pd.notna(row.iloc[1]) and str(row.iloc[1]) != 'nan':
                    name = str(row.iloc[1]).strip()
                    designation = str(row.iloc[2]).strip() if pd.notna(row.iloc[2]) else ''
                    
                    if name and not name[0].isdigit():
                        faculty = {
                            'si_no': row.iloc[0] if pd.notna(row.iloc[0]) else None,
                            'name': name,
                            'designation': designation
                        }
                        faculty_data.append(faculty)
        
        print(f"Total faculty found: {len(faculty_data)}")
        return faculty_data, course_assignments
        
    except Exception as e:
        print(f"Error extracting faculty from Excel: {e}")
        import traceback
        traceback.print_exc()
        return [], []

def extract_section_tables_from_docx(filepath):
    """Extract tables from Word document for 3rd, 5th, and 7th semesters"""
    try:
        doc = Document(filepath)
        
        tables_data = {
            '3rd_semester': [],
            '5th_semester': [],
            '7th_semester': []
        }
        
        print(f"Found {len(doc.tables)} tables in document")
        
        for table_idx, table in enumerate(doc.tables):
            if len(table.rows) < 4:
                continue
            
            first_row_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
            
            if 'SECOND YEAR' in first_row_text or 'COORINDATOR' in first_row_text:
                print(f"Table {table_idx}: Found 3rd semester table")
                
                for row_idx in range(3, len(table.rows)):
                    row = table.rows[row_idx]
                    cells = row.cells
                    
                    if len(cells) >= 8:
                        course_name = cells[2].text.strip().replace('\n', ' ').replace('  ', ' ')
                        
                        course_data = {
                            'sl_no': cells[0].text.strip(),
                            'course_code': cells[1].text.strip(),
                            'course_name': course_name,
                            'section_A': cells[3].text.strip(),
                            'section_B': cells[4].text.strip(),
                            'section_C': cells[5].text.strip(),
                            'section_D': cells[6].text.strip(),
                            'coordinator': cells[7].text.strip() if len(cells) > 7 else ''
                        }
                        
                        if course_data['course_code'] and course_data['course_code'] not in ['', 'Course Code']:
                            tables_data['3rd_semester'].append(course_data)
                            print(f"  Found course: {course_data['course_code']} - {course_data['course_name'][:50]}")
            
            elif 'THIRD YEAR' in first_row_text or 'COORDINATOR' in first_row_text:
                print(f"Table {table_idx}: Found 5th semester table")
                
                for row_idx in range(3, len(table.rows)):
                    row = table.rows[row_idx]
                    cells = row.cells
                    
                    if len(cells) >= 8:
                        course_name = cells[2].text.strip().replace('\n', ' ').replace('  ', ' ')
                        
                        course_data = {
                            'sl_no': cells[0].text.strip(),
                            'course_code': cells[1].text.strip(),
                            'course_name': course_name,
                            'section_A': cells[3].text.strip(),
                            'section_B': cells[4].text.strip(),
                            'section_C': cells[5].text.strip(),
                            'section_D': cells[6].text.strip(),
                            'coordinator': cells[7].text.strip() if len(cells) > 7 else ''
                        }
                        
                        if course_data['course_code'] and course_data['course_code'] not in ['', 'Course Code']:
                            tables_data['5th_semester'].append(course_data)
                            print(f"  Found course: {course_data['course_code']} - {course_data['course_name'][:50]}")
        
        return tables_data
        
    except Exception as e:
        print(f"Error extracting tables from Word: {e}")
        import traceback
        traceback.print_exc()
        return {}

def extract_course_ratios(course_name):
    """Extract lecture:tutorial:lab ratios from course name"""
    ratio_pattern = r'\((\d+):(\d+):(\d+)\)'
    match = re.search(ratio_pattern, course_name)
    
    if match:
        lectures = int(match.group(1))
        tutorials = int(match.group(2))
        labs = int(match.group(3))
        return lectures, tutorials, labs
    else:
        if 'LAB' in course_name.upper() or 'LABORATORY' in course_name.upper():
            return 0, 0, 1
        elif 'TUT' in course_name.upper() or 'TUTORIAL' in course_name.upper():
            return 0, 1, 0
        else:
            return 3, 0, 0

def generate_faculty_mapping(faculty_data, course_assignments, section_tables):
    """Generate faculty ID mapping with max units based on designation"""
    
    faculty_mapping = {}
    
    designation_limits = {
        'Head & Professor': 18,
        'Professor': 18,
        'Associate Professor': 24,
        'Assistant Professor': 28,
        'Senior Lecturer': 28
    }
    
    for faculty in faculty_data:
        name = faculty['name'].strip()
        faculty_id = re.sub(r'[^\w\s]', '', name).replace(' ', '_').lower()
        
        designation = faculty['designation']
        max_units = 18
        for key in designation_limits:
            if key in designation:
                max_units = designation_limits[key]
                break
        
        faculty_mapping[faculty_id] = {
            'id': faculty_id,
            'name': name,
            'original_name': name,
            'designation': designation,
            'max_units': max_units
        }
    
    return faculty_mapping

def normalize_name(name):
    """Normalize faculty name for comparison"""
    if not name:
        return ""
    name = name.replace('.', '').lower().strip()
    name = ' '.join(name.split())
    return name

def build_course_mapping(section_tables, faculty_mapping):
    """Build mapping of section -> course_code -> faculty_id"""
    
    course_mapping = {
        '3A': {}, '3B': {}, '3C': {}, '3D': {},
        '5A': {}, '5B': {}, '5C': {}, '5D': {},
        '7A': {}, '7B': {}
    }
    
    # Tutorial display name mapping
    tutorial_names = {
        '24CS35': 'DMS TUT',
        '23CS53': 'DBS TUT',
    }
    
    # Complete initials mapping
    initials_mapping = {
        'MRC': 'Dr.Monica R Mundada',
        'ASB': 'Dr.Akshata S Bhayyar',
        'AKK': 'Akshata Kamath',
        'SB': 'Dr.Sushma B',
        'SJ': 'Dr.J.Sangeetha',
        'NSB': 'Nandini S B',
        'NSV': 'Nandini S B',
        'SRR': 'Dr.Rajarajeswari S',
        'MA': 'Mamatha A',
        'US': 'Uzma Sulthana',
        'DPK': 'Pradeep kumar D',
        'JGR': 'Dr.J Geetha',
        'SCS': 'Soumya C S',
        'SS': 'Dr.Sangeetha V',
        'APK': 'Dr.A Parkavi',
        'VGS': 'Veena GS',
        'MG': 'Mallegowda M',
        'CP': 'Chandrika Prasad',
        'PK': 'Priya K',
        'DN': 'Darshana A Naik',
        'TNR': 'Dr.T.N R. Kumar',
        'DRB': 'Dr.Dayananda R B',
        'SV': 'Dr.Sangeetha V',
        'BG': 'Brunda',
        'VCD': 'Vishwachetan D',
        'GIS': 'Dr.Ganeshayya I Shidaganti',
        'MRM': 'Dr.Monica R Mundada',
        'SSC': 'Dr.Shilpa S Chaudhari',
        'RCA': 'Dr. China Appala Naidu',
        'JSM': 'Jamuna S Murthy',
    }
    
    # Direct mapping for problematic initials
    direct_faculty_mapping = {
        'AKK': 'akshata_kamath',
        'SB': 'drsushma_b',
        'PK': 'priya_k',
        'BG': 'brunda',
        'ASB': 'drakshata_s_bhayyar',
        'NSB': 'nandini_s_b',
        'SCS': 'soumya_c_s',
        'SSC': 'drshilpa_s_chaudhari',
    }
    
    # Create normalized lookup
    normalized_faculty_map = {}
    for faculty_id, faculty_info in faculty_mapping.items():
        normalized_name = normalize_name(faculty_info['name'])
        normalized_faculty_map[normalized_name] = faculty_id
        normalized_faculty_map[faculty_id] = faculty_id
    
    def find_faculty_id(short_name):
        if not short_name:
            return None
        
        if '/' in short_name:
            short_name = short_name.split('/')[0]
        
        if short_name in direct_faculty_mapping:
            return direct_faculty_mapping[short_name]
        
        if short_name in initials_mapping:
            full_name = initials_mapping[short_name]
            normalized_full = normalize_name(full_name)
            if normalized_full in normalized_faculty_map:
                return normalized_faculty_map[normalized_full]
        
        cleaned_short = short_name.replace('.', '').lower().strip()
        for normalized, faculty_id in normalized_faculty_map.items():
            if cleaned_short in normalized:
                return faculty_id
        
        for faculty_id, faculty_info in faculty_mapping.items():
            faculty_name_normalized = normalize_name(faculty_info['name'])
            if cleaned_short in faculty_name_normalized:
                return faculty_id
        
        print(f"    WARNING: Could not find faculty for '{short_name}'")
        return None
    
    print("\n" + "="*60)
    print("BUILDING COURSE MAPPING")
    print("="*60)
    
    # Create a list of ALL 3rd semester courses including the missing one
    all_3rd_courses = list(section_tables.get('3rd_semester', []))
    
    missing_24cs32 = True
    for c in all_3rd_courses:
        if c['course_code'] == '24CS32':
            missing_24cs32 = False
            break
    
    if missing_24cs32:
        print("\n  ➕ MANUALLY ADDING MISSING COURSE: 24CS32")
        missing_course = {
            'sl_no': '1',
            'course_code': '24CS32',
            'course_name': 'Digital Design and Computer Organization (3:0:1)',
            'section_A': 'MRC',
            'section_B': 'ASB',
            'section_C': 'MRC',
            'section_D': 'ASB',
            'coordinator': 'ASB'
        }
        all_3rd_courses.append(missing_course)
    
    # Process 3rd semester courses
    print("\n📚 3RD SEMESTER COURSES:")
    for course in all_3rd_courses:
        course_code = course['course_code']
        course_name = course['course_name']
        
        lectures, tutorials, labs = extract_course_ratios(course_name)
        
        print(f"\n  Course: {course_code} (L:{lectures}, T:{tutorials}, Lab:{labs})")
        
        for section in ['A', 'B', 'C', 'D']:
            section_key = f"3{section}"
            faculty_short = course.get(f'section_{section}', '')
            
            if faculty_short:
                faculty_id = find_faculty_id(faculty_short)
                if faculty_id:
                    # For courses with tutorials, modify the course name to remove tutorial ratio
                    if tutorials > 0 and course_code in tutorial_names:
                        # Create modified course name without tutorial
                        modified_name = re.sub(r'\(\d+:\d+:\d+\)', '(3:0:0)', course_name)
                        if '2:1:0' in course_name:
                            modified_name = course_name.replace('(2:1:0)', '(2:0:0)')
                        course_mapping[section_key][course_code] = {
                            'faculty_id': faculty_id,
                            'course_name': modified_name,
                            'lectures': lectures,
                            'tutorials': 0,
                            'labs': labs
                        }
                        print(f"    {section_key}: {course_code} -> {faculty_id} (Modified to L:{lectures}, T:0, Lab:{labs})")
                        
                        # Create separate tutorial course
                        tutorial_code = f"{course_code}_TUT"
                        tutorial_display_name = tutorial_names[course_code]
                        course_mapping[section_key][tutorial_code] = {
                            'faculty_id': faculty_id,
                            'course_name': f"{tutorial_display_name} (0:1:0)",
                            'lectures': 0,
                            'tutorials': 1,
                            'labs': 0
                        }
                        print(f"    {section_key}: {tutorial_code} (TUTORIAL) -> {faculty_id}")
                    else:
                        # Regular course (no tutorial separation needed)
                        course_mapping[section_key][course_code] = {
                            'faculty_id': faculty_id,
                            'course_name': course_name,
                            'lectures': lectures,
                            'tutorials': tutorials,
                            'labs': labs
                        }
                        print(f"    {section_key}: {course_code} -> {faculty_id} (L:{lectures}, T:{tutorials}, Lab:{labs})")
                else:
                    print(f"    {section_key}: {faculty_short} -> NOT FOUND")
            else:
                print(f"    {section_key}: NO FACULTY")
    
    # Process 5th semester courses
    all_5th_courses = list(section_tables.get('5th_semester', []))
    
    missing_23cs51 = True
    for c in all_5th_courses:
        if c['course_code'] == '23CS51':
            missing_23cs51 = False
            break
    
    if missing_23cs51:
        print("\n  ➕ MANUALLY ADDING MISSING COURSE: 23CS51")
        missing_course_5th = {
            'sl_no': '1',
            'course_code': '23CS51',
            'course_name': 'Software Engineering & Modeling (3:0:0)',
            'section_A': 'DN',
            'section_B': 'TNR',
            'section_C': 'DPK',
            'section_D': 'APK',
            'coordinator': 'DN'
        }
        all_5th_courses.append(missing_course_5th)
    
    print("\n📚 5TH SEMESTER COURSES:")
    for course in all_5th_courses:
        course_code = course['course_code']
        course_name = course['course_name']
        
        lectures, tutorials, labs = extract_course_ratios(course_name)
        
        print(f"\n  Course: {course_code} (L:{lectures}, T:{tutorials}, Lab:{labs})")
        
        for section in ['A', 'B', 'C', 'D']:
            section_key = f"5{section}"
            faculty_short = course.get(f'section_{section}', '')
            
            # Handle elective course section assignments
            if faculty_short and ('Section-' in faculty_short or 'Section ' in faculty_short):
                if course_code == '23CSE551':
                    if section == 'A':
                        faculty_short = 'NSB'
                    elif section == 'B':
                        faculty_short = 'SV'
                    elif section == 'C':
                        faculty_short = 'SCS'
                    else:
                        faculty_short = ''
            
            if faculty_short:
                faculty_id = find_faculty_id(faculty_short)
                if faculty_id:
                    # For courses with tutorials, modify the course name to remove tutorial ratio
                    if tutorials > 0 and course_code in tutorial_names:
                        # Create modified course name without tutorial
                        modified_name = re.sub(r'\(\d+:\d+:\d+\)', '(3:0:0)', course_name)
                        if '2:1:0' in course_name:
                            modified_name = course_name.replace('(2:1:0)', '(2:0:0)')
                        course_mapping[section_key][course_code] = {
                            'faculty_id': faculty_id,
                            'course_name': modified_name,
                            'lectures': lectures,
                            'tutorials': 0,
                            'labs': labs
                        }
                        print(f"    {section_key}: {course_code} -> {faculty_id} (Modified to L:{lectures}, T:0, Lab:{labs})")
                        
                        # Create separate tutorial course
                        tutorial_code = f"{course_code}_TUT"
                        tutorial_display_name = tutorial_names[course_code]
                        course_mapping[section_key][tutorial_code] = {
                            'faculty_id': faculty_id,
                            'course_name': f"{tutorial_display_name} (0:1:0)",
                            'lectures': 0,
                            'tutorials': 1,
                            'labs': 0
                        }
                        print(f"    {section_key}: {tutorial_code} (TUTORIAL) -> {faculty_id}")
                    else:
                        # Regular course (no tutorial separation needed)
                        course_mapping[section_key][course_code] = {
                            'faculty_id': faculty_id,
                            'course_name': course_name,
                            'lectures': lectures,
                            'tutorials': tutorials,
                            'labs': labs
                        }
                        print(f"    {section_key}: {course_code} -> {faculty_id} (L:{lectures}, T:{tutorials}, Lab:{labs})")
                else:
                    print(f"    {section_key}: {faculty_short} -> NOT FOUND")
            else:
                print(f"    {section_key}: NO FACULTY")
    
    return course_mapping

def build_course_ratios(section_tables, course_mapping):
    """Build course ratios from course names"""
    course_ratios = {}
    
    all_courses = {}
    for section, courses in course_mapping.items():
        for course_code, course_info in courses.items():
            if course_code not in all_courses:
                all_courses[course_code] = course_info['course_name']
    
    print("\n📈 Course Ratios:")
    for course_code, course_name in all_courses.items():
        lectures, tutorials, labs = extract_course_ratios(course_name)
        course_ratios[course_code] = {
            'course_name': course_name,
            'lectures_per_week': lectures,
            'tutorials_per_week': tutorials,
            'labs_per_week': labs,
            'total_slots_per_week': lectures + tutorials + labs
        }
        print(f"   {course_code}: {lectures}L + {tutorials}T + {labs}Lab = {lectures + tutorials + labs} slots")
    
    return course_ratios

def main():
    excel_file = DATA_DIR / 'AY 2025_2026 ODD SUBJECT ALLOTMENT_JULY 2025.xlsx'
    docx_file = DATA_DIR / 'Section wise Subject Allotment.docx'
    
    print("="*60)
    print("DEPARTMENT TIMETABLE DATA EXTRACTOR")
    print("="*60)
    
    print("\n📁 Extracting data from Excel file...")
    faculty_data, course_assignments = extract_faculty_from_excel(excel_file)
    print(f"   Found {len(faculty_data)} faculty members")
    
    print("\n📄 Extracting data from Word document...")
    section_tables = extract_section_tables_from_docx(docx_file)
    
    print(f"\n📊 Extraction Results:")
    print(f"   3rd semester: {len(section_tables.get('3rd_semester', []))} courses extracted")
    print(f"   5th semester: {len(section_tables.get('5th_semester', []))} courses extracted")
    
    # Check and report missing courses
    expected_3rd = ['24CS32', '24CS33', '24CS34', '24CS35', '24CSL36', '24CSL37', '24UHV38', '24CSAEC310']
    actual_3rd = [c['course_code'] for c in section_tables.get('3rd_semester', [])]
    missing_3rd = [c for c in expected_3rd if c not in actual_3rd]
    if missing_3rd:
        print(f"\n⚠️  Missing 3rd semester courses: {missing_3rd}")
        print(f"   These will be ADDED MANUALLY")
    
    expected_5th = ['23CS51', '23CS52', '23CS53', '23CS54', '23CSE551', '23CSE5511', '23CSL56', '23CSL57', '23AL58', '23CSAEC510']
    actual_5th = [c['course_code'] for c in section_tables.get('5th_semester', [])]
    missing_5th = [c for c in expected_5th if c not in actual_5th]
    if missing_5th:
        print(f"⚠️  Missing 5th semester courses: {missing_5th}")
        print(f"   These will be ADDED MANUALLY")
    
    print("\n👥 Generating faculty mapping...")
    faculty_mapping = generate_faculty_mapping(faculty_data, course_assignments, section_tables)
    
    print("\n🗺️ Building course mapping...")
    course_mapping = build_course_mapping(section_tables, faculty_mapping)
    
    print("\n📈 Building course ratios...")
    course_ratios = build_course_ratios(section_tables, course_mapping)
    
    print("\n💾 Saving to JSON files...")
    
    with open(OUTPUT_DIR / 'faculty_data.json', 'w', encoding='utf-8') as f:
        json.dump(faculty_mapping, f, indent=2, ensure_ascii=False)
    print("[OK] faculty_data.json created")
    
    with open(OUTPUT_DIR / 'course_mapping.json', 'w', encoding='utf-8') as f:
        json.dump(course_mapping, f, indent=2, ensure_ascii=False)
    print("[OK] course_mapping.json created")
    
    with open(OUTPUT_DIR / 'course_ratios.json', 'w', encoding='utf-8') as f:
        json.dump(course_ratios, f, indent=2, ensure_ascii=False)
    print("[OK] course_ratios.json created")
    
    # Final Summary
    print("\n" + "="*60)
    print("EXTRACTION SUMMARY")
    print("="*60)
    print(f"Total faculty members: {len(faculty_mapping)}")
    print(f"Total sections: {len(course_mapping)}")
    
    total_courses = 0
    for section, courses in course_mapping.items():
        if section.startswith('3') or section.startswith('5'):
            print(f"  {section}: {len(courses)} courses")
            total_courses += len(courses)
    print(f"\nTotal course assignments: {total_courses}")
    print(f"Total unique courses: {len(course_ratios)}")
    
    # Verify critical courses
    print("\n" + "="*60)
    print("CRITICAL COURSE VERIFICATION")
    print("="*60)
    
    for section in ['3A', '3B', '3C', '3D']:
        if '24CS32' in course_mapping.get(section, {}):
            print(f"✅ {section}: 24CS32 present")
        else:
            print(f"❌ {section}: 24CS32 MISSING!")
        
        if '24CS35_TUT' in course_mapping.get(section, {}):
            print(f"✅ {section}: 24CS35_TUT (DMS TUT) present")
        else:
            print(f"❌ {section}: 24CS35_TUT (DMS TUT) MISSING!")
    
    for section in ['5A', '5B', '5C', '5D']:
        if '23CS51' in course_mapping.get(section, {}):
            print(f"✅ {section}: 23CS51 present")
        else:
            print(f"❌ {section}: 23CS51 MISSING!")
        
        if '23CS53_TUT' in course_mapping.get(section, {}):
            print(f"✅ {section}: 23CS53_TUT (DBS TUT) present")
        else:
            print(f"❌ {section}: 23CS53_TUT (DBS TUT) MISSING!")
    
    print("\n" + "="*60)
    print("✅ EXTRACTION COMPLETE!")
    print("="*60)

if __name__ == "__main__":
    main()